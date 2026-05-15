"""Document Generator, Risk Assessment & Case Roadmap API routes.

Document generation uses a 3-stage Google ADK agent pipeline:
  1. SituationAnalyzerAgent  – detects country / region / legal domain / laws
  2. TemplateResearchAgent   – finds the correct format for that jurisdiction
  3. DocumentDrafterAgent    – produces the final plain-text legal document
"""

import io
import json
import os
import re
from datetime import datetime
from enum import Enum

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.services.gemma_service import gemma_service
from app.agents.docgen_agents import generate_document_pipeline
from app.core.security import rate_limit_dependency, sanitize_input
from app.core.logging import logger

router = APIRouter(prefix="/api/docgen", tags=["docgen"])


# ── Schemas ──────────────────────────────────────────────────────────────

class DocType(str, Enum):
    RTI_APPLICATION = "rti_application"
    LEGAL_NOTICE = "legal_notice"
    COMPLAINT_LETTER = "complaint_letter"
    BAIL_APPLICATION = "bail_application"
    CONSUMER_COMPLAINT = "consumer_complaint"
    FIR_COMPLAINT = "fir_complaint"
    CUSTOM = "custom"


DOC_TYPE_LABELS = {
    "rti_application": "RTI Application",
    "legal_notice": "Legal Notice",
    "complaint_letter": "Complaint Letter",
    "bail_application": "Bail Application",
    "consumer_complaint": "Consumer Complaint",
    "fir_complaint": "FIR Complaint",
    "custom": "Custom Legal Document",
}


class DocGenRequest(BaseModel):
    doc_type: DocType
    context: str = Field(..., min_length=10, max_length=5000)
    language: str = Field(default="en")
    user_name: str = Field(default="[Your Name]", max_length=200)
    user_address: str = Field(default="[Your Address]", max_length=500)
    opposite_party: str = Field(default="[Opposite Party]", max_length=500)
    custom_doc_title: str = Field(default="", max_length=300)


class RiskRequest(BaseModel):
    context: str = Field(..., min_length=10, max_length=5000)
    language: str = Field(default="en")


class RoadmapRequest(BaseModel):
    context: str = Field(..., min_length=10, max_length=5000)
    legal_domain: str = Field(default="")
    language: str = Field(default="en")


# ── Document Generation Endpoint (3-Agent ADK Pipeline) ──────────────

@router.post("/generate", dependencies=[Depends(rate_limit_dependency)])
async def generate_document(req: DocGenRequest):
    """Generate a ready-to-use legal document via the 3-agent ADK pipeline."""
    context = sanitize_input(req.context)

    # For custom doc type, use the user-provided title
    if req.doc_type.value == "custom" and req.custom_doc_title:
        doc_type_label = sanitize_input(req.custom_doc_title)
    else:
        doc_type_label = DOC_TYPE_LABELS.get(req.doc_type.value, req.doc_type.value)

    try:
        result = await generate_document_pipeline(
            doc_type=req.doc_type.value if req.doc_type.value != "custom" else doc_type_label,
            doc_type_label=doc_type_label,
            context=context,
            language=req.language,
            user_name=sanitize_input(req.user_name),
            user_address=sanitize_input(req.user_address),
            opposite_party=sanitize_input(req.opposite_party),
        )
        return {
            "doc_type": req.doc_type.value,
            "doc_type_label": doc_type_label,
            "content": result["document"],
            "situation_analysis": result.get("situation_analysis"),
            "template_used": result.get("template_used"),
            "language": req.language,
            "generated_at": datetime.utcnow().isoformat(),
        }
    except Exception as e:
        logger.error(f"DocGen pipeline error: {e}")
        return {
            "doc_type": req.doc_type.value,
            "content": "Unable to generate the document. Please try again.",
            "language": req.language,
            "error": True,
        }


# ── Document PDF Download (fpdf2 with Unicode) ──────────────────────

class DocDownloadRequest(BaseModel):
    title: str = Field(default="Legal Document", max_length=200)
    content: str = Field(..., min_length=1)
    language: str = Field(default="en")


@router.post("/pdf")
async def download_doc_pdf(req: DocDownloadRequest):
    """Download a generated document as a PDF with full Unicode support."""
    try:
        pdf_bytes = _build_unicode_pdf(req.title, req.content, req.language)
    except Exception as e:
        import traceback
        logger.error(f"DocGen PDF error: {e}\n{traceback.format_exc()}")
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")

    doc_slug = re.sub(r'[^a-z0-9]+', '_', req.title.lower()).strip('_') or 'document'
    timestamp = int(datetime.utcnow().timestamp() * 1000)
    filename = f"dharma_nyaya_{doc_slug}_{timestamp}.pdf"

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Document Word Download (.docx) ─────────────────────────────────

@router.post("/word")
async def download_doc_word(req: DocDownloadRequest):
    """Download a generated document as a Word .docx file."""
    try:
        docx_bytes = _build_word_doc(req.title, req.content, req.language)
    except Exception as e:
        import traceback
        logger.error(f"DocGen Word error: {e}\n{traceback.format_exc()}")
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Word generation failed: {e}")

    doc_slug = re.sub(r'[^a-z0-9]+', '_', req.title.lower()).strip('_') or 'document'
    timestamp = int(datetime.utcnow().timestamp() * 1000)
    filename = f"dharma_nyaya_{doc_slug}_{timestamp}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Risk Assessment Endpoint ─────────────────────────────────────────────

@router.post("/risk-assess", dependencies=[Depends(rate_limit_dependency)])
async def assess_risk(req: RiskRequest):
    """Analyze a legal situation and return a risk assessment with score."""
    context = sanitize_input(req.context)

    prompt = f"""Analyze the following legal situation and provide a risk assessment.

You MUST respond with ONLY a valid JSON object (no markdown, no code fences, no commentary). Use this exact structure:

{{
  "risk_score": <number 1-100>,
  "risk_level": "<low|medium|high|critical>",
  "success_probability": <number 1-100>,
  "summary": "<2-3 sentence risk summary>",
  "risk_factors": ["<risk factor 1>", "<risk factor 2>", ...],
  "strengths": ["<strength 1>", "<strength 2>", ...],
  "recommendations": ["<recommendation 1>", "<recommendation 2>", ...],
  "time_estimate": "<estimated timeline for resolution>",
  "estimated_cost": "<rough cost estimate range>",
  "references": [
    {{"title": "<official source name for the detected jurisdiction>", "uri": "<exact real URL from the official legal site of that country>"}}
  ]
}}

JURISDICTION-AWARE REFERENCE RULES:
- Detect jurisdiction from language of the query, country/law names mentioned, or explicit location.
- India (default): indiankanoon.org, indiacode.nic.in, sci.gov.in, legislative.gov.in, labour.gov.in, consumeraffairs.nic.in, nalsa.gov.in, rtionline.gov.in
- USA: law.cornell.edu, uscode.house.gov, ftc.gov, dol.gov, consumerfinance.gov, justia.com
- UK: legislation.gov.uk, gov.uk, citizensadvice.org.uk, bailii.org
- Ukraine: zakon.rada.gov.ua, minjust.gov.ua, court.gov.ua
- EU: eur-lex.europa.eu, curia.europa.eu
- Australia: legislation.gov.au, austlii.edu.au, accc.gov.au
- Canada: laws-lois.justice.gc.ca, canlii.org
- Other: official national parliament and government legal portals of that country.
- NEVER mix jurisdictions. Include 2-4 real URLs from the detected jurisdiction ONLY.

Risk scoring guide:
- 1-25: Low risk (strong case, favorable outcome likely)
- 26-50: Medium risk (reasonable case, needs preparation)
- 51-75: High risk (significant challenges, uncertain outcome)
- 76-100: Critical risk (very difficult, major obstacles)

Legal situation: {context}"""

    try:
        result = await gemma_service.generate_text(
            prompt,
            language=req.language,
            system_instruction=(
                "You are a legal risk analyst. Respond with ONLY valid JSON. "
                "No markdown, no code blocks, no extra text. Be realistic and honest."
            ),
        )

        # Parse JSON from response
        assessment = _extract_json_object(result)
        if not assessment:
            assessment = {
                "risk_score": 50,
                "risk_level": "medium",
                "success_probability": 50,
                "summary": result[:300],
                "risk_factors": [],
                "strengths": [],
                "recommendations": [],
                "time_estimate": "Unknown",
                "estimated_cost": "Varies",
            }

        # Clamp values
        assessment["risk_score"] = max(1, min(100, int(assessment.get("risk_score", 50))))
        assessment["success_probability"] = max(1, min(100, int(assessment.get("success_probability", 50))))
        assessment["language"] = req.language
        assessment["references"] = [
            {"title": r.get("title", ""), "uri": r.get("uri", "")}
            for r in assessment.get("references", [])
            if str(r.get("uri", "")).startswith("http")
        ]
        return assessment

    except Exception as e:
        logger.error(f"Risk assessment error: {e}")
        return {
            "risk_score": 50,
            "risk_level": "medium",
            "success_probability": 50,
            "summary": "Unable to complete risk assessment. Please try again.",
            "risk_factors": [],
            "strengths": [],
            "recommendations": [],
            "language": req.language,
            "error": True,
        }


# ── Case Roadmap Endpoint ────────────────────────────────────────────────

@router.post("/roadmap", dependencies=[Depends(rate_limit_dependency)])
async def generate_roadmap(req: RoadmapRequest):
    """Generate a step-by-step legal process roadmap for the situation."""
    context = sanitize_input(req.context)

    prompt = f"""Analyze the following legal situation and generate a step-by-step legal roadmap.

You MUST respond with ONLY a valid JSON object (no markdown, no code fences). Use this exact structure:

{{
  "title": "<Short title for the case roadmap>",
  "total_steps": <number>,
  "estimated_duration": "<total estimated time>",
  "steps": [
    {{
      "step": 1,
      "title": "<step title>",
      "description": "<what to do in this step>",
      "timeline": "<estimated time for this step>",
      "documents_needed": ["<doc 1>", "<doc 2>"],
      "tips": "<helpful tip for this step>",
      "icon": "<emoji for this step>"
    }}
  ],
  "important_notes": ["<note 1>", "<note 2>"],
  "legal_provisions": ["<relevant law/section 1>", "<relevant law/section 2>"],
  "references": [
    {{"title": "<official source name for the detected jurisdiction>", "uri": "<exact real URL from the official legal site of that country>"}}
  ]
}}

JURISDICTION-AWARE REFERENCE RULES:
- Detect jurisdiction from language of the query, country/law names mentioned, or explicit location.
- India (default): indiankanoon.org, indiacode.nic.in, sci.gov.in, legislative.gov.in, labour.gov.in, consumeraffairs.nic.in, nalsa.gov.in, rtionline.gov.in
- USA: law.cornell.edu, uscode.house.gov, ftc.gov, dol.gov, consumerfinance.gov, justia.com
- UK: legislation.gov.uk, gov.uk, citizensadvice.org.uk, bailii.org
- Ukraine: zakon.rada.gov.ua, minjust.gov.ua, court.gov.ua
- EU: eur-lex.europa.eu, curia.europa.eu
- Australia: legislation.gov.au, austlii.edu.au, hcourt.gov.au
- Canada: laws-lois.justice.gc.ca, canlii.org
- Other: official national parliament and government legal portals of that country.
- NEVER mix jurisdictions. Include 2-4 real URLs from the detected jurisdiction ONLY.
Provide 5-8 concrete, actionable steps. Be specific to the detected jurisdiction's legal system.
{f'Legal domain: {req.legal_domain}' if req.legal_domain else ''}

Legal situation: {context}"""

    try:
        result = await gemma_service.generate_text(
            prompt,
            language=req.language,
            system_instruction=(
                "You are a legal process expert covering multiple jurisdictions worldwide. "
                "Respond with ONLY valid JSON. No markdown, no code blocks. "
                "Provide practical, step-by-step guidance."
            ),
        )

        roadmap = _extract_json_object(result)
        if not roadmap or "steps" not in roadmap:
            roadmap = {
                "title": "Legal Process Roadmap",
                "steps": [{"step": 1, "title": "Consult a Lawyer", "description": result[:500],
                           "timeline": "Immediately", "documents_needed": [], "tips": "", "icon": "📋"}],
                "important_notes": [],
                "legal_provisions": [],
            }

        roadmap["language"] = req.language
        roadmap["references"] = [
            {"title": r.get("title", ""), "uri": r.get("uri", "")}
            for r in roadmap.get("references", [])
            if str(r.get("uri", "")).startswith("http")
        ]
        return roadmap

    except Exception as e:
        logger.error(f"Roadmap generation error: {e}")
        return {
            "title": "Legal Process Roadmap",
            "steps": [],
            "important_notes": ["Unable to generate roadmap. Please try again."],
            "legal_provisions": [],
            "language": req.language,
            "error": True,
        }


# ── Helper Functions ─────────────────────────────────────────────────────

def _extract_json_object(text: str) -> dict | None:
    """Extract a JSON object from potentially noisy LLM output."""
    # Try direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Strip markdown fences
    cleaned = re.sub(r'```(?:json)?\s*', '', text)
    cleaned = re.sub(r'```\s*$', '', cleaned)
    try:
        return json.loads(cleaned.strip())
    except json.JSONDecodeError:
        pass

    # Find first { to last }
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass

    return None


def _build_unicode_pdf(title: str, content: str, language: str = "en") -> bytes:
    """Build a PDF with full Unicode support using fpdf2 and NirmalaUI / NotoSans."""
    from fpdf import FPDF

    # Font directory in the project
    font_dir = os.path.join(os.path.dirname(__file__), "..", "..", "..", "fonts")
    font_dir = os.path.abspath(font_dir)

    font_regular = None
    font_bold = None

    # Prefer NirmalaUI (covers Devanagari, Bengali, Tamil, Telugu, Kannada, Latin)
    nirmala_r = os.path.join(font_dir, "NirmalaUI-Regular.ttf")
    nirmala_b = os.path.join(font_dir, "NirmalaUI-Bold.ttf")
    noto_r = os.path.join(font_dir, "NotoSans-Regular.ttf")

    if os.path.isfile(nirmala_r):
        font_regular = nirmala_r
        font_bold = nirmala_b if os.path.isfile(nirmala_b) else nirmala_r
    elif os.path.isfile(noto_r):
        font_regular = noto_r
        font_bold = noto_r

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.add_page()

    # Register Unicode font
    if font_regular:
        pdf.add_font("DocFont", "", font_regular, uni=True)
        pdf.add_font("DocFont", "B", font_bold, uni=True)
        font_family = "DocFont"
    else:
        font_family = "Helvetica"

    # ── Title ──
    pdf.set_font(font_family, "B", 14)
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.multi_cell(effective_w, 8, title, align="C")
    pdf.set_font(font_family, "", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(effective_w, 5, f"Generated by DHARMA-NYAYA AI  |  {datetime.utcnow().strftime('%d %B %Y')}", align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    # Separator line
    pdf.set_draw_color(180, 180, 180)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    # ── Body ──
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin
    for line in content.split("\n"):
        stripped = line.strip()

        if not stripped:
            pdf.ln(4)
            continue

        # Detect UPPERCASE headings (the drafter agent uses them)
        is_heading = (
            stripped == stripped.upper()
            and len(stripped) > 3
            and any(c.isalpha() for c in stripped)
            and not stripped.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '('))
        )

        if is_heading:
            pdf.ln(3)
            pdf.set_font(font_family, "B", 11)
            pdf.multi_cell(effective_w, 6, stripped)
            pdf.ln(2)
            continue

        # Regular text
        pdf.set_font(font_family, "", 10)
        pdf.multi_cell(effective_w, 5, stripped)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def _build_word_doc(title: str, content: str, language: str = "en") -> bytes:
    """Build a Word .docx document from the generated legal text."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(55, 48, 163)  # indigo

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"Generated by DHARMA-NYAYA AI  |  {datetime.utcnow().strftime('%d %B %Y')}"
    )
    sub_run.font.size = Pt(8)
    sub_run.font.color.rgb = RGBColor(120, 120, 120)

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("─" * 60)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(180, 180, 180)

    # Body content
    for line in content.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()  # blank line
            continue

        # Detect UPPERCASE headings
        is_heading = (
            stripped == stripped.upper()
            and len(stripped) > 3
            and any(c.isalpha() for c in stripped)
            and not stripped.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '('))
        )

        if is_heading:
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(stripped)
            h_run.bold = True
            h_run.font.size = Pt(12)
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _wrap(text: str, width: int) -> list[str]:
    """Simple word wrap."""
    words = text.split()
    lines = []
    current = ""
    for w in words:
        if len(current) + len(w) + 1 > width:
            if current:
                lines.append(current)
            current = w
        else:
            current = f"{current} {w}" if current else w
    if current:
        lines.append(current)
    return lines or [""]
